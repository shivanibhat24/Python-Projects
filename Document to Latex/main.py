#!/usr/bin/env python3
"""
Enhanced Document to LaTeX Converter

This script converts various document formats (Word, PDF, plain text, Markdown, or HTML)
to LaTeX format. The script detects the input format and applies appropriate conversion rules.

Requirements:
    - python-docx (for Word documents)
    - PyPDF2 (for PDF documents)
    - python-mammoth (for improved Word document handling)

Installation:
    pip install python-docx PyPDF2 mammoth

Usage:
    python doc_to_latex.py input_file [output_file.tex] --title "Document Title" --author "Author Name"
    
If no output file is specified, the output will be written to stdout.
"""

import sys
import os
import re
import argparse
from html.parser import HTMLParser
import mimetypes
import subprocess
from pathlib import Path

# Try importing specialized libraries
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False


class HTMLToLatexParser(HTMLParser):
    """HTML parser that converts HTML to LaTeX format."""
    
    def __init__(self):
        super().__init__()
        self.result = []
        self.in_list = False
        self.list_items = []
        self.in_table = False
        self.in_row = False
        self.row_cells = []
        self.table_rows = []
        self.in_header = False
        self.image_count = 0
        
    def handle_starttag(self, tag, attrs):
        if tag == 'h1':
            self.result.append("\\section{")
            self.in_header = True
        elif tag == 'h2':
            self.result.append("\\subsection{")
            self.in_header = True
        elif tag == 'h3':
            self.result.append("\\subsubsection{")
            self.in_header = True
        elif tag == 'h4':
            self.result.append("\\paragraph{")
            self.in_header = True
        elif tag == 'h5' or tag == 'h6':
            self.result.append("\\subparagraph{")
            self.in_header = True
        elif tag == 'p':
            self.result.append("\n\n")
        elif tag == 'em' or tag == 'i':
            self.result.append("\\emph{")
        elif tag == 'strong' or tag == 'b':
            self.result.append("\\textbf{")
        elif tag == 'u':
            self.result.append("\\underline{")
        elif tag == 'ul':
            self.in_list = True
            self.list_items = []
        elif tag == 'ol':
            self.in_list = True
            self.list_items = []
            self.result.append("\n\\begin{enumerate}\n")
        elif tag == 'li' and self.in_list:
            self.result.append("\\item ")
        elif tag == 'a':
            href = next((attr[1] for attr in attrs if attr[0] == 'href'), None)
            if href:
                self.result.append(f"\\href{{{href}}}{{{")
        elif tag == 'table':
            self.in_table = True
            self.table_rows = []
        elif tag == 'tr' and self.in_table:
            self.in_row = True
            self.row_cells = []
        elif tag == 'img':
            src = next((attr[1] for attr in attrs if attr[0] == 'src'), None)
            alt = next((attr[1] for attr in attrs if attr[0] == 'alt'), f"Image {self.image_count}")
            self.image_count += 1
            if src:
                filename = os.path.basename(src)
                self.result.append(f"\n\\begin{{figure}}[h]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{{{filename}}}}}\n\\caption{{{alt}}}\n\\end{{figure}}\n")
        
    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.result.append("}")
            self.in_header = False
        elif tag in ['em', 'i', 'strong', 'b', 'u']:
            self.result.append("}")
        elif tag == 'ul':
            self.in_list = False
            self.result.append("\n\\begin{itemize}\n" + "".join(self.list_items) + "\\end{itemize}\n")
            self.list_items = []
        elif tag == 'ol':
            self.in_list = False
            self.result.append("\n" + "".join(self.list_items) + "\\end{enumerate}\n")
            self.list_items = []
        elif tag == 'a':
            self.result.append("}")
        elif tag == 'table':
            # Format collected table data
            if self.table_rows:
                # Determine number of columns from the first row
                num_cols = max(len(row) for row in self.table_rows)
                col_spec = "|" + "c|" * num_cols
                
                table_str = f"\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
                for i, row in enumerate(self.table_rows):
                    # Pad rows with empty cells if needed
                    padded_row = row + [""] * (num_cols - len(row))
                    table_str += " & ".join(padded_row) + " \\\\ \\hline\n"
                table_str += "\\end{tabular}"
                
                self.result.append(f"\n\\begin{{center}}\n{table_str}\n\\end{{center}}\n")
            self.in_table = False
        elif tag == 'tr':
            if self.row_cells:
                self.table_rows.append(self.row_cells)
            self.in_row = False
        
    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.in_list and self.result and self.result[-1] == "\\item ":
                self.list_items.append(f"\\item {self.escape_latex(data)}\n")
            elif self.in_row and self.in_table:
                self.row_cells.append(self.escape_latex(data))
            else:
                self.result.append(self.escape_latex(data))
    
    def escape_latex(self, text):
        """Escape special LaTeX characters."""
        chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '^': '\\textasciicircum{}',
            '\\': '\\textbackslash{}',
        }
        # Don't escape already escaped characters
        result = ""
        i = 0
        while i < len(text):
            if text[i:i+2] in ['\\\\']: 
                result += text[i:i+2]
                i += 2
            elif text[i] in chars:
                result += chars[text[i]]
                i += 1
            else:
                result += text[i]
                i += 1
        return result
    
    def get_result(self):
        return "".join(self.result)


def detect_format(content, file_extension=None):
    """Attempt to detect the format of the document based on content and extension."""
    if file_extension:
        ext = file_extension.lower()
        if ext in ['.docx', '.doc']:
            return 'word'
        elif ext == '.pdf':
            return 'pdf'
        elif ext in ['.html', '.htm']:
            return 'html'
        elif ext in ['.md', '.markdown']:
            return 'markdown'
        elif ext == '.txt':
            return 'plaintext'
    
    # Check for HTML
    if re.search(r'<(?:html|body|p|h[1-6]|div|span|table|ul|ol|li)\b', content, re.IGNORECASE):
        return 'html'
    
    # Check for Markdown headers, lists, or emphasis
    md_patterns = [
        r'^#{1,6}\s+.+$',  # Headers
        r'^\*\s+.+$',      # Unordered lists
        r'^\d+\.\s+.+$',   # Ordered lists
        r'\*\*.+\*\*',     # Bold
        r'_.+_',           # Italic
    ]
    for pattern in md_patterns:
        if re.search(pattern, content, re.MULTILINE):
            return 'markdown'
    
    # Default to plain text
    return 'plaintext'


def convert_markdown_to_latex(content):
    """Convert Markdown to LaTeX."""
    # Headers
    content = re.sub(r'^# (.+)$', r'\\section{\1}', content, flags=re.MULTILINE)
    content = re.sub(r'^## (.+)$', r'\\subsection{\1}', content, flags=re.MULTILINE)
    content = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', content, flags=re.MULTILINE)
    content = re.sub(r'^#### (.+)$', r'\\paragraph{\1}', content, flags=re.MULTILINE)
    content = re.sub(r'^##### (.+)$', r'\\subparagraph{\1}', content, flags=re.MULTILINE)
    
    # Bold and italic
    content = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', content)
    content = re.sub(r'__(.+?)__', r'\\textbf{\1}', content)
    content = re.sub(r'\*(.+?)\*', r'\\emph{\1}', content)
    content = re.sub(r'_(.+?)_', r'\\emph{\1}', content)
    
    # Links
    content = re.sub(r'\[(.+?)\]\((.+?)\)', r'\\href{\2}{\1}', content)
    
    # Images
    content = re.sub(r'!\[(.+?)\]\((.+?)\)', r'\\begin{figure}[h]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{\2}\n\\caption{\1}\n\\end{figure}', content)
    
    # Code blocks
    content = re.sub(r'```(.+?)```', r'\\begin{verbatim}\1\\end{verbatim}', content, flags=re.DOTALL)
    content = re.sub(r'`(.+?)`', r'\\texttt{\1}', content)
    
    # Lists
    # This is simplified; a more robust approach would use a parser
    # Unordered lists
    list_pattern = r'^(\s*)\*\s+(.+)$'
    list_matches = re.findall(list_pattern, content, re.MULTILINE)
    if list_matches:
        list_blocks = []
        for indent, item in list_matches:
            list_blocks.append(f"\\item {item}")
        
        content = re.sub(list_pattern, '', content, flags=re.MULTILINE)
        content += "\n\\begin{itemize}\n" + "\n".join(list_blocks) + "\n\\end{itemize}\n"
    
    # Ordered lists
    ordered_list_pattern = r'^(\s*)\d+\.\s+(.+)$'
    ordered_list_matches = re.findall(ordered_list_pattern, content, re.MULTILINE)
    if ordered_list_matches:
        list_blocks = []
        for indent, item in ordered_list_matches:
            list_blocks.append(f"\\item {item}")
        
        content = re.sub(ordered_list_pattern, '', content, flags=re.MULTILINE)
        content += "\n\\begin{enumerate}\n" + "\n".join(list_blocks) + "\n\\end{enumerate}\n"
    
    # Escape special LaTeX characters not in already processed structures
    special_chars = ['&', '%', '$', '#', '_', '{', '}', '~', '^']
    for char in special_chars:
        # Don't replace already escaped characters
        escaped_char = f'\\{char}'
        i = 0
        new_content = ""
        while i < len(content):
            if content[i:i+len(escaped_char)] == escaped_char:
                new_content += escaped_char
                i += len(escaped_char)
            elif content[i] == char:
                new_content += escaped_char
                i += 1
            else:
                new_content += content[i]
                i += 1
        content = new_content
    
    return content


def convert_plaintext_to_latex(content):
    """Convert plain text to LaTeX, with minimal formatting."""
    # Escape special characters
    special_chars = ['&', '%', '$', '#', '_', '{', '}', '~', '^', '\\']
    for char in special_chars:
        content = content.replace(char, f'\\{char}')
    
    # Convert paragraphs (blank lines between text)
    paragraphs = re.split(r'\n\s*\n', content)
    content = '\n\n'.join(paragraphs)
    
    return content


def convert_html_to_latex(content):
    """Convert HTML to LaTeX using the HTMLParser."""
    parser = HTMLToLatexParser()
    parser.feed(content)
    return parser.get_result()


def generate_latex_document(content, title="Document", author="Author"):
    """Generate a complete LaTeX document with the converted content."""
    latex_template = f"""\\documentclass{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{hyperref}}
\\usepackage{{graphicx}}
\\usepackage{{booktabs}}
\\usepackage{{multirow}}
\\usepackage{{tabularx}}
\\usepackage{{float}}
\\usepackage{{listings}}
\\usepackage{{color}}
\\usepackage{{xcolor}}

\\definecolor{{dkgreen}}{{rgb}}{{0,0.6,0}}
\\definecolor{{gray}}{{rgb}}{{0.5,0.5,0.5}}
\\definecolor{{mauve}}{{rgb}}{{0.58,0,0.82}}

\\lstset{{
  frame=tb,
  language=Python,
  aboveskip=3mm,
  belowskip=3mm,
  showstringspaces=false,
  columns=flexible,
  basicstyle=\\small\\ttfamily,
  numbers=none,
  numberstyle=\\tiny\\color{{gray}},
  keywordstyle=\\color{{blue}},
  commentstyle=\\color{{dkgreen}},
  stringstyle=\\color{{mauve}},
  breaklines=true,
  breakatwhitespace=true,
  tabsize=3
}}

\\title{{{title}}}
\\author{{{author}}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle

{content}

\\end{{document}}
"""
    return latex_template


def extract_text_from_docx(docx_path):
    """Extract text from a Word document using python-docx."""
    if not DOCX_AVAILABLE:
        print("Warning: python-docx is not installed. Basic text extraction will be used.")
        return extract_text_from_docx_basic(docx_path)
    
    doc = docx.Document(docx_path)
    full_text = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Process tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            table_text.append(" | ".join(row_text))
        full_text.append("\n".join(table_text))
    
    return "\n\n".join(full_text)


def extract_text_from_docx_basic(docx_path):
    """A basic method to extract text from a Word document without python-docx."""
    # Try to use unzip if available
    temp_dir = Path("temp_docx_extraction")
    try:
        if not temp_dir.exists():
            temp_dir.mkdir()
        
        subprocess.run(["unzip", "-q", docx_path, "word/document.xml", "-d", str(temp_dir)])
        
        with open(temp_dir / "word" / "document.xml", "r", encoding="utf-8") as f:
            xml_content = f.read()
        
        # Very basic extraction of text between <w:t> tags
        text_parts = re.findall(r'<w:t.*?>(.*?)</w:t>', xml_content, re.DOTALL)
        full_text = " ".join(text_parts)
        
        # Clean up
        import shutil
        shutil.rmtree(temp_dir)
        
        return full_text
    except Exception as e:
        print(f"Error extracting text from Word document: {e}")
        if temp_dir.exists():
            import shutil
            shutil.rmtree(temp_dir)
        return f"Error extracting text from {docx_path}: {str(e)}"


def extract_text_with_mammoth(docx_path):
    """Extract text from a Word document using Mammoth, which preserves more formatting."""
    if not MAMMOTH_AVAILABLE:
        print("Warning: mammoth is not installed. Basic extraction will be used.")
        if DOCX_AVAILABLE:
            return extract_text_from_docx(docx_path)
        else:
            return extract_text_from_docx_basic(docx_path)
    
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
        
        # Convert the HTML to LaTeX
        return convert_html_to_latex(html)


def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF document."""
    if not PDF_AVAILABLE:
        print("Warning: PyPDF2 is not installed. PDF conversion will not work.")
        return f"Error: PyPDF2 is required to extract text from {pdf_path}"
    
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text() + "\n\n"
        return text


def main():
    parser = argparse.ArgumentParser(description='Convert documents to LaTeX format')
    parser.add_argument('input_file', help='Input document file')
    parser.add_argument('output_file', nargs='?', help='Output LaTeX file (optional)')
    parser.add_argument('--title', default="Document", help='Document title')
    parser.add_argument('--author', default="Author", help='Document author')
    args = parser.parse_args()
    
    # Check if required libraries are installed
    if not DOCX_AVAILABLE:
        print("Note: python-docx is not installed. Word document conversion may be limited.")
    if not PDF_AVAILABLE:
        print("Note: PyPDF2 is not installed. PDF conversion will not work.")
    if not MAMMOTH_AVAILABLE:
        print("Note: mammoth is not installed. Word document conversion will use basic methods.")
    
    # Get file extension
    file_extension = os.path.splitext(args.input_file)[1]
    
    # Process based on file type
    if file_extension.lower() in ['.docx', '.doc']:
        if MAMMOTH_AVAILABLE:
            content = extract_text_with_mammoth(args.input_file)
        else:
            content = extract_text_from_docx(args.input_file)
            content = convert_plaintext_to_latex(content)
    elif file_extension.lower() == '.pdf':
        content = extract_text_from_pdf(args.input_file)
        content = convert_plaintext_to_latex(content)
    else:
        # Read input file for other formats
        try:
            with open(args.input_file, 'r', encoding='utf-8') as file:
                content = file.read()
        except Exception as e:
            print(f"Error reading input file: {e}", file=sys.stderr)
            sys.exit(1)
        
        # Detect format and convert
        doc_format = detect_format(content, file_extension)
        
        if doc_format == 'html':
            content = convert_html_to_latex(content)
        elif doc_format == 'markdown':
            content = convert_markdown_to_latex(content)
        else:  # plaintext
            content = convert_plaintext_to_latex(content)
    
    # Generate complete LaTeX document
    latex_document = generate_latex_document(content, args.title, args.author)
    
    # Output result
    if args.output_file:
        try:
            with open(args.output_file, 'w', encoding='utf-8') as file:
                file.write(latex_document)
            print(f"Conversion successful. Output written to {args.output_file}")
        except Exception as e:
            print(f"Error writing output file: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        print(latex_document)


if __name__ == "__main__":
    main()
