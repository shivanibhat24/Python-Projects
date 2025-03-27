import os
import re
import json
import spacy
import docx
import PyPDF2
import pandas as pd
import numpy as np
from typing import List, Dict, Any
from dataclasses import dataclass, asdict
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

class RequirementExtractor:
    def __init__(self):
        # Load spaCy English model
        self.nlp = spacy.load("en_core_web_sm")
        
        # Knowledge base for standard requirements and best practices
        self.knowledge_base = self._load_knowledge_base()
        
        # TF-IDF Vectorizer for semantic similarity
        self.vectorizer = TfidfVectorizer(stop_words='english')

    def _load_knowledge_base(self) -> Dict[str, Any]:
        """
        Load predefined knowledge base for requirements
        In a real-world scenario, this would be more comprehensive
        """
        return {
            "software_standards": [
                "Follow SOLID principles",
                "Implement proper error handling",
                "Ensure scalability",
                "Maintain code modularity"
            ],
            "common_non_functional_requirements": [
                "Performance",
                "Security",
                "Usability",
                "Reliability",
                "Compatibility"
            ]
        }

    def extract_requirements(self, text: str) -> Dict[str, List[str]]:
        """
        Extract functional and non-functional requirements
        """
        doc = self.nlp(text)
        
        # Extract potential requirements using linguistic patterns
        functional_reqs = []
        non_functional_reqs = []
        
        # Basic requirement extraction using linguistic patterns
        for sent in doc.sents:
            sent_text = sent.text.lower()
            
            # Detect requirement linguistic patterns
            requirement_patterns = [
                r'shall\s',
                r'should\s',
                r'must\s',
                r'will\s',
                r'need\s+to',
                r'required\s+to'
            ]
            
            for pattern in requirement_patterns:
                if re.search(pattern, sent_text):
                    # Classify requirements
                    if any(nf in sent_text for nf in self.knowledge_base['common_non_functional_requirements']):
                        non_functional_reqs.append(sent.text)
                    else:
                        functional_reqs.append(sent.text)
        
        return {
            "functional_requirements": functional_reqs,
            "non_functional_requirements": non_functional_reqs
        }

    def prioritize_requirements(self, requirements: List[str]) -> List[Dict[str, Any]]:
        """
        Prioritize requirements using MoSCoW method
        """
        def _classify_priority(req: str) -> str:
            req_lower = req.lower()
            if any(word in req_lower for word in ['must', 'critical', 'essential']):
                return 'Must Have'
            elif any(word in req_lower for word in ['should', 'important']):
                return 'Should Have'
            elif any(word in req_lower for word in ['could', 'optional', 'nice to have']):
                return 'Could Have'
            else:
                return 'Won\'t Have'
        
        return [
            {
                "requirement": req,
                "priority": _classify_priority(req),
                "complexity_score": len(req.split())  # Simple complexity metric
            }
            for req in requirements
        ]

    def document_parser(self, file_path: str) -> str:
        """
        Parse different document types
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    return ' '.join([page.extract_text() for page in reader.pages])
            
            elif file_ext == '.docx':
                doc = docx.Document(file_path)
                return ' '.join([para.text for para in doc.paragraphs])
            
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        
        except Exception as e:
            print(f"Error parsing document: {e}")
            return ""

    def generate_requirement_document(self, requirements: Dict[str, List[str]], output_path: str):
        """
        Generate a standardized requirement document
        """
        doc = docx.Document()
        
        # Title
        doc.add_heading('Requirements Specification Document', 0)
        
        # Functional Requirements
        doc.add_heading('Functional Requirements', level=1)
        for req in requirements.get('functional_requirements', []):
            doc.add_paragraph(req, style='List Bullet')
        
        # Non-Functional Requirements
        doc.add_heading('Non-Functional Requirements', level=1)
        for req in requirements.get('non_functional_requirements', []):
            doc.add_paragraph(req, style='List Bullet')
        
        doc.save(output_path)

    def generate_jira_backlog(self, prioritized_requirements: List[Dict[str, Any]], output_path: str):
        """
        Generate Excel sheet for Jira backlog
        """
        df = pd.DataFrame(prioritized_requirements)
        df.to_excel(output_path, index=False)

class RequirementsManager:
    def __init__(self):
        self.extractor = RequirementExtractor()
        self.version_history = {}

    def process_document(self, file_path: str):
        """
        Main method to process requirement documents
        """
        # Extract text from document
        document_text = self.extractor.document_parser(file_path)
        
        # Extract requirements
        requirements = self.extractor.extract_requirements(document_text)
        
        # Prioritize requirements
        prioritized_reqs = {
            "functional": self.extractor.prioritize_requirements(requirements['functional_requirements']),
            "non_functional": self.extractor.prioritize_requirements(requirements['non_functional_requirements'])
        }
        
        # Generate artifacts
        version_id = self._generate_version_id(file_path)
        self.version_history[version_id] = requirements
        
        # Generate Word document
        word_output = f"requirements_{version_id}.docx"
        self.extractor.generate_requirement_document(requirements, word_output)
        
        # Generate Jira backlog
        excel_output = f"jira_backlog_{version_id}.xlsx"
        self.extractor.generate_jira_backlog(
            prioritized_reqs['functional'] + prioritized_reqs['non_functional'], 
            excel_output
        )
        
        return {
            "version_id": version_id,
            "requirements": requirements,
            "prioritized_requirements": prioritized_reqs,
            "word_document": word_output,
            "excel_backlog": excel_output
        }

    def _generate_version_id(self, file_path: str) -> str:
        """
        Generate unique version identifier
        """
        import hashlib
        from datetime import datetime
        
        # Create a hash based on file content and timestamp
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{timestamp}_{file_hash[:8]}"

def main():
    # Example usage
    manager = RequirementsManager()
    result = manager.process_document("sample_requirements.docx")
    
    print("Requirements Extracted:")
    print(json.dumps(result, indent=2))

if __name__ == "__main__":
    main()
